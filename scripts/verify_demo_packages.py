#!/usr/bin/env python3
"""Validate and render the three fictional textbook demo-data packages.

This script writes only disposable QA previews below ``tmp/demo-data-build/qa``.
It does not rebuild or change the downloadable ZIP files.
"""

from __future__ import annotations

import concurrent.futures
import csv
import hashlib
import io
import json
import os
import re
import shutil
import stat
import subprocess
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path, PurePosixPath

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from pypdf import PdfReader


REPO = Path(__file__).resolve().parents[1]
BUILD = REPO / "tmp" / "demo-data-build"
PUBLIC = REPO / "public" / "downloads" / "demo-data"
QA = BUILD / "qa"
NOTICE = "練習用・すべて架空・外部送信禁止"

INDUSTRIES = {
    "salon": {
        "root": "美容室デモデータ",
        "zip": "salon-demo-data-full-v1.zip",
        "label": "ひだまり美容室",
        "company": "株式会社ひだまりデモサロン（架空）",
        "business": "美容室2店舗の運営、予約、施術、物販、集客",
        "sheets": 44,
        "dataRows": 130571,
    },
    "construction": {
        "root": "建設業デモデータ",
        "zip": "construction-demo-data-full-v1.zip",
        "label": "青空デモ建設",
        "company": "株式会社青空デモ建設（架空）",
        "business": "住宅改修、店舗改修、小規模新築、外構工事",
        "sheets": 45,
        "dataRows": 99230,
    },
    "realestate": {
        "root": "不動産会社デモデータ",
        "zip": "real-estate-demo-data-full-v1.zip",
        "label": "まちかどデモ不動産",
        "company": "株式会社まちかどデモ不動産（架空）",
        "business": "賃貸仲介、売買仲介、賃貸管理、修繕受付",
        "sheets": 46,
        "dataRows": 100050,
    },
}

STARTERS = [
    "01メール.txt",
    "02新聞.txt",
    "03画像.txt",
    "04見積.txt",
    "05資料.txt",
    "06HP.txt",
    "07商談.txt",
    "08朝新聞.txt",
    "09アプリ.txt",
    "10AI安全.txt",
]

LESSON_REFERENCES = ["11会社.pdf", "12料金.pdf", "13組織.pdf", "14AIルール.docx"]


def fail(message: str) -> None:
    raise RuntimeError(message)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def sha256_bytes(payload: bytes) -> str:
    return hashlib.sha256(payload).hexdigest()


def parse_checksums(path: Path) -> dict[str, str]:
    lines = path.read_text(encoding="utf-8").rstrip("\n").splitlines()
    if len(lines) != len(INDUSTRIES):
        fail(f"{path.name} must contain exactly {len(INDUSTRIES)} lines")
    result: dict[str, str] = {}
    for number, line in enumerate(lines, start=1):
        match = re.fullmatch(r"([0-9a-f]{64})  ([^/\\\r\n]+\.zip)", line)
        if not match:
            fail(f"Invalid {path.name} line {number}: {line!r}")
        checksum, file_name = match.groups()
        if file_name in result:
            fail(f"Duplicate filename in {path.name}: {file_name}")
        result[file_name] = checksum
    expected = [details["zip"] for details in INDUSTRIES.values()]
    if list(result) != expected:
        fail(f"Unexpected filename/order in {path.name}: {list(result)}")
    return result


def build_constant(path: Path, name: str) -> str:
    source = path.read_text(encoding="utf-8")
    match = re.search(
        rf"(?:const\s+)?{re.escape(name)}\s*=\s*['\"]([^'\"]+)['\"]",
        source,
    )
    if not match:
        fail(f"{name} not found in {path.relative_to(REPO)}")
    return match.group(1)


def validate_build_versions(catalog: dict) -> None:
    version = catalog.get("version")
    updated_at = catalog.get("updatedAt")
    if not re.fullmatch(r"\d+\.\d+\.\d+", str(version)):
        fail(f"Invalid catalog version: {version!r}")
    if not re.fullmatch(r"\d{4}-\d{2}-\d{2}", str(updated_at)):
        fail(f"Invalid catalog updatedAt: {updated_at!r}")
    for script in (REPO / "scripts/build_demo_data.mjs", REPO / "scripts/build_demo_documents.py"):
        if build_constant(script, "VERSION") != version:
            fail(f"VERSION differs between catalog.json and {script.name}")
        if build_constant(script, "GENERATED_AT") != updated_at:
            fail(f"GENERATED_AT differs between catalog.json and {script.name}")


def source_and_zip_hashes(
    package_root: Path,
    zip_path: Path,
    root_name: str,
) -> tuple[dict[str, str], dict[str, str]]:
    source_files = {
        path.relative_to(package_root).as_posix(): sha256(path)
        for path in sorted(package_root.rglob("*"))
        if path.is_file()
    }
    with zipfile.ZipFile(zip_path) as archive:
        zip_files: dict[str, str] = {}
        prefix = f"{root_name}/"
        for info in archive.infolist():
            if info.is_dir():
                continue
            if not info.filename.startswith(prefix):
                fail(f"Unexpected ZIP root: {info.filename}")
            relative = info.filename.removeprefix(prefix)
            if relative in zip_files:
                fail(f"Duplicate path in {zip_path.name}: {relative}")
            zip_files[relative] = sha256_bytes(archive.read(info))
    if source_files.keys() != zip_files.keys():
        missing = sorted(source_files.keys() - zip_files.keys())
        extra = sorted(zip_files.keys() - source_files.keys())
        fail(f"Source/ZIP path mismatch in {zip_path.name}: missing={missing}, extra={extra}")
    mismatched = [name for name in source_files if source_files[name] != zip_files[name]]
    if mismatched:
        fail(f"Source/ZIP SHA-256 mismatch in {zip_path.name}: {mismatched}")
    return source_files, zip_files


def parse_hash_lines(text: str, label: str) -> dict[str, str]:
    result: dict[str, str] = {}
    for number, line in enumerate(text.rstrip("\n").splitlines(), start=1):
        match = re.fullmatch(r"([0-9a-f]{64})  (.+)", line)
        if not match:
            fail(f"Invalid {label} line {number}: {line!r}")
        checksum, relative = match.groups()
        if relative in result:
            fail(f"Duplicate filename in {label}: {relative}")
        result[relative] = checksum
    return result


def validate_package_manifest(package_root: Path, source_hashes: dict[str, str]) -> str:
    manifest_path = package_root / "manifest.csv"
    with manifest_path.open(encoding="utf-8-sig", newline="") as source:
        reader = csv.DictReader(source)
        if reader.fieldnames != ["path", "format", "bytes", "sha256", "notice"]:
            fail(f"Unexpected manifest columns: {reader.fieldnames}")
        rows = list(reader)
    expected_manifest_paths = set(source_hashes) - {"manifest.csv", "SHA256SUMS.txt"}
    actual_manifest_paths = {row["path"] for row in rows}
    if actual_manifest_paths != expected_manifest_paths or len(rows) != len(actual_manifest_paths):
        fail(f"manifest.csv path mismatch in {package_root.name}")
    for row in rows:
        file_path = package_root / PurePosixPath(row["path"])
        if row["notice"] != NOTICE:
            fail(f"Manifest notice mismatch: {row['path']}")
        if row["format"] != file_path.suffix.lower().lstrip("."):
            fail(f"Manifest format mismatch: {row['path']}")
        if int(row["bytes"]) != file_path.stat().st_size:
            fail(f"Manifest byte count mismatch: {row['path']}")
        if row["sha256"] != source_hashes[row["path"]]:
            fail(f"Manifest SHA-256 mismatch: {row['path']}")

    internal_hashes = parse_hash_lines(
        (package_root / "SHA256SUMS.txt").read_text(encoding="utf-8"),
        f"{package_root.name}/SHA256SUMS.txt",
    )
    expected_hash_paths = set(source_hashes) - {"SHA256SUMS.txt"}
    if internal_hashes.keys() != expected_hash_paths:
        fail(f"Internal SHA256SUMS path mismatch in {package_root.name}")
    for relative, checksum in internal_hashes.items():
        if checksum != source_hashes[relative]:
            fail(f"Internal SHA-256 mismatch: {package_root.name}/{relative}")
    return source_hashes["manifest.csv"]


def xlsx_metrics(paths: list[Path]) -> tuple[int, int]:
    worksheet_pattern = re.compile(r"xl/worksheets/sheet\d+\.xml")
    sheet_count = 0
    data_rows = 0
    namespace = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"
    for path in paths:
        with zipfile.ZipFile(path) as archive:
            if archive.testzip() is not None:
                fail(f"XLSX CRC error: {path}")
            worksheets = [name for name in archive.namelist() if worksheet_pattern.fullmatch(name)]
            if not worksheets:
                fail(f"XLSX has no worksheets: {path}")
            sheet_count += len(worksheets)
            for worksheet in worksheets:
                root = ET.fromstring(archive.read(worksheet))
                data_rows += sum(
                    1
                    for row in root.findall(f".//{namespace}row")
                    if int(row.get("r", "0")) >= 5
                )
    return sheet_count, data_rows


def validate_zip(path: Path, root_name: str) -> tuple[int, int, int]:
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            fail(f"ZIP CRC error: {path.name}")
        names = [item.filename for item in archive.infolist()]
        for item in archive.infolist():
            pure = PurePosixPath(item.filename)
            if pure.is_absolute() or ".." in pure.parts:
                fail(f"Unsafe ZIP path: {item.filename}")
            mode = item.external_attr >> 16
            if stat.S_ISLNK(mode):
                fail(f"Symlink found in ZIP: {item.filename}")
            if ".inspect.ndjson" in item.filename or "/Users/" in item.filename:
                fail(f"Private/build file leaked into ZIP: {item.filename}")
            if pure.parts and pure.parts[0] != root_name:
                fail(f"Unexpected ZIP root: {item.filename}")

        starter_prefix = f"{root_name}/課題/"
        actual_starters = sorted(
            PurePosixPath(name).name
            for name in names
            if name.startswith(starter_prefix) and name.endswith(".txt")
        )
        if actual_starters != sorted(STARTERS):
            fail(f"Starter memo mismatch in {path.name}: {actual_starters}")
        for file_name in LESSON_REFERENCES:
            if f"{starter_prefix}{file_name}" not in names:
                fail(f"Lesson reference missing in {path.name}: {file_name}")
        if f"{root_name}/完成/ここに完成品を入れます.txt" not in names:
            fail(f"Completion folder marker missing in {path.name}")

        return (
            sum(name.endswith(".xlsx") for name in names),
            sum(name.endswith(".docx") for name in names),
            sum(name.endswith(".pdf") for name in names),
        )


def xlsx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        archive.testzip()
        payload = b"".join(
            archive.read(name)
            for name in archive.namelist()
            if name.endswith(".xml")
        )
    return payload.decode("utf-8", errors="ignore")


def xlsx_bytes_text(payload: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        if archive.testzip() is not None:
            fail("XLSX CRC error in downloadable ZIP")
        xml_payload = b"".join(
            archive.read(name)
            for name in archive.namelist()
            if name.endswith(".xml")
        )
    return xml_payload.decode("utf-8", errors="ignore")


def xlsx_has_notice(path: Path) -> bool:
    return "架空" in xlsx_text(path)


def docx_text(path: Path) -> str:
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            fail(f"DOCX CRC error: {path}")
    document = Document(path)
    pieces = [paragraph.text for paragraph in document.paragraphs]
    pieces.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(pieces)


def docx_bytes_text(payload: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(payload)) as archive:
        if archive.testzip() is not None:
            fail("DOCX CRC error in downloadable ZIP")
    document = Document(io.BytesIO(payload))
    pieces = [paragraph.text for paragraph in document.paragraphs]
    pieces.extend(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    return "\n".join(pieces)


def pdf_text(path: Path) -> str:
    reader = PdfReader(path)
    if not reader.pages:
        fail(f"PDF has no pages: {path}")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def pdf_bytes_text(payload: bytes) -> str:
    reader = PdfReader(io.BytesIO(payload))
    if not reader.pages:
        fail("PDF has no pages in downloadable ZIP")
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def zip_readable_text(path: Path) -> tuple[str, str]:
    with zipfile.ZipFile(path) as archive:
        if archive.testzip() is not None:
            fail(f"ZIP CRC error: {path.name}")
        readable: list[str] = []
        readme_text = ""
        for name in archive.namelist():
            suffix = PurePosixPath(name).suffix.lower()
            if suffix not in {".txt", ".csv", ".json", ".xlsx", ".docx", ".pdf"}:
                continue
            payload = archive.read(name)
            if suffix in {".txt", ".csv", ".json"}:
                text = payload.decode("utf-8", errors="ignore")
            elif suffix == ".xlsx":
                text = xlsx_bytes_text(payload)
            elif suffix == ".docx":
                text = docx_bytes_text(payload)
            else:
                text = pdf_bytes_text(payload)
            readable.append(text)
            if name.endswith("/README_最初に読んでください.txt"):
                readme_text = text
    if not readme_text:
        fail(f"Starter README missing in {path.name}")
    return readme_text, "\n".join(readable)


def pdf_has_embedded_font(path: Path) -> bool:
    reader = PdfReader(path)
    for page in reader.pages:
        fonts = page.get("/Resources", {}).get("/Font", {})
        for font_ref in fonts.values():
            font = font_ref.get_object()
            descriptors = []
            if font.get("/FontDescriptor"):
                descriptors.append(font["/FontDescriptor"].get_object())
            for descendant_ref in font.get("/DescendantFonts", []):
                descendant = descendant_ref.get_object()
                if descendant.get("/FontDescriptor"):
                    descriptors.append(descendant["/FontDescriptor"].get_object())
            if any(
                any(key in descriptor for key in ("/FontFile", "/FontFile2", "/FontFile3"))
                for descriptor in descriptors
            ):
                return True
    return False


def find_docx_renderer() -> Path:
    override = os.environ.get("DOCX_RENDERER")
    if override:
        return Path(override)
    candidates = list(
        Path.home().glob(
            ".cache/codex-runtimes/codex-primary-runtime/plugins/"
            "openai-primary-runtime/plugins/documents/skills/documents/"
            "render_docx.py"
        )
    )
    if not candidates:
        fail("render_docx.py was not found. Set DOCX_RENDERER.")
    return candidates[0]


def find_docx_renderer_python(renderer: Path) -> Path:
    if override := os.environ.get("DOCX_RENDERER_PYTHON"):
        candidate = Path(override)
        if not candidate.is_file():
            fail(f"DOCX_RENDERER_PYTHON does not exist: {candidate}")
        return candidate
    for parent in renderer.parents:
        if parent.name == "codex-primary-runtime":
            candidate = parent / "dependencies/python/bin/python3"
            if candidate.is_file():
                return candidate
    return Path(sys.executable)


def run_checked(
    command: list[str],
    label: str,
    *,
    env: dict[str, str] | None = None,
) -> subprocess.CompletedProcess[str]:
    result = subprocess.run(
        command,
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        env=env,
    )
    if result.returncode != 0:
        fail(
            f"{label} failed with exit {result.returncode}\n"
            f"STDOUT:\n{result.stdout.strip()}\n"
            f"STDERR:\n{result.stderr.strip()}"
        )
    return result


def render_docx(path: Path, renderer: Path, renderer_python: Path) -> list[Path]:
    relative = path.relative_to(BUILD)
    output = QA / "docx" / relative.with_suffix("")
    output.mkdir(parents=True, exist_ok=True)
    for old_page in output.glob("page-*.png"):
        old_page.unlink()
    emitted_pdf = output / f"{path.stem}.pdf"
    emitted_pdf.unlink(missing_ok=True)
    render_env = dict(os.environ)
    if sys.platform == "darwin" and Path("/private/tmp").is_dir():
        render_env.update({"TMPDIR": "/private/tmp", "TEMP": "/private/tmp", "TMP": "/private/tmp"})
    macos_fontconfig = Path("/opt/homebrew/etc/fonts/fonts.conf")
    if macos_fontconfig.exists():
        render_env["FONTCONFIG_FILE"] = str(macos_fontconfig)
        render_env["FONTCONFIG_PATH"] = str(macos_fontconfig.parent)
    run_checked(
        [
            str(renderer_python),
            str(renderer),
            str(path),
            "--output_dir",
            str(output),
            "--width",
            "900",
            "--height",
            "1200",
            "--emit_pdf",
        ],
        f"DOCX render for {path}",
        env=render_env,
    )
    pages = sorted(output.glob("*.png"))
    if not pages:
        fail(f"DOCX did not render: {path}")
    if not emitted_pdf.is_file() or emitted_pdf.stat().st_size == 0:
        fail(f"DOCX renderer did not emit a PDF: {path}")
    emitted_text = pdf_text(emitted_pdf)
    if NOTICE not in emitted_text:
        fail(f"Rendered DOCX PDF lost Japanese notice text: {path}")
    if "\ufffd" in emitted_text:
        fail(f"Rendered DOCX PDF contains replacement characters: {path}")
    return pages


def render_pdf(path: Path, pdftoppm: str) -> list[Path]:
    relative = path.relative_to(BUILD)
    output = QA / "pdf" / relative.with_suffix("")
    output.parent.mkdir(parents=True, exist_ok=True)
    for old_page in output.parent.glob(f"{output.name}-*.png"):
        old_page.unlink()
    run_checked(
        [
            pdftoppm,
            "-png",
            "-scale-to-x",
            "700",
            "-scale-to-y",
            "-1",
            str(path),
            str(output),
        ],
        f"PDF render for {path}",
    )
    rendered = sorted(output.parent.glob(f"{output.name}-*.png"))
    if not rendered:
        fail(f"PDF did not render: {path}")
    return rendered


def contact_sheets(images: list[Path], name: str, per_sheet: int = 20) -> list[Path]:
    destination = QA / "contact-sheets"
    destination.mkdir(parents=True, exist_ok=True)
    for old_sheet in destination.glob(f"{name}-*.jpg"):
        old_sheet.unlink()
    created: list[Path] = []
    columns = 4
    rows = per_sheet // columns
    cell_w, cell_h = 360, 300
    font = ImageFont.load_default()
    for start in range(0, len(images), per_sheet):
        group = images[start : start + per_sheet]
        canvas = Image.new("RGB", (columns * cell_w, rows * cell_h), "white")
        draw = ImageDraw.Draw(canvas)
        for index, image_path in enumerate(group):
            with Image.open(image_path) as source:
                preview = source.convert("RGB")
                preview.thumbnail((cell_w - 20, cell_h - 52), Image.Resampling.LANCZOS)
                x = (index % columns) * cell_w + (cell_w - preview.width) // 2
                y = (index // columns) * cell_h + 8
                canvas.paste(preview, (x, y))
            label = image_path.stem[:42]
            draw.text(
                ((index % columns) * cell_w + 10, (index // columns) * cell_h + cell_h - 36),
                label,
                fill="black",
                font=font,
            )
        page = start // per_sheet + 1
        output = destination / f"{name}-{page:02d}.jpg"
        canvas.save(output, quality=86)
        created.append(output)
    return created


def main() -> int:
    catalog = json.loads((PUBLIC / "catalog.json").read_text(encoding="utf-8"))
    if set(catalog) != {"version", "updatedAt", "packages"}:
        fail(f"Unexpected catalog root keys: {sorted(catalog)}")
    validate_build_versions(catalog)
    packages = catalog.get("packages")
    if not isinstance(packages, list) or len(packages) != len(INDUSTRIES):
        fail(f"Catalog must contain exactly {len(INDUSTRIES)} packages")
    if [item.get("industry") for item in packages] != list(INDUSTRIES):
        fail("Catalog package order/industries do not match the three supported packages")
    catalog_by_industry = {item["industry"]: item for item in packages}
    checksums = parse_checksums(PUBLIC / "SHA256SUMS.txt")

    expected_public_zips = sorted(details["zip"] for details in INDUSTRIES.values())
    actual_public_zips = sorted(path.name for path in PUBLIC.glob("*.zip") if path.is_file())
    if actual_public_zips != expected_public_zips:
        fail(f"Unexpected public ZIP set: {actual_public_zips}")

    all_docx: list[Path] = []
    all_pdf: list[Path] = []
    expected_xlsx_previews: set[Path] = set()
    summary: dict[str, object] = {"industries": {}, "renders": {}}
    total_workbooks = 0
    total_sheets = 0
    total_data_rows = 0

    expected_catalog_keys = {
        "industry",
        "label",
        "company",
        "business",
        "file",
        "url",
        "bytes",
        "sha256",
        "version",
        "updatedAt",
        "workbooks",
        "sheets",
        "dataRows",
        "docx",
        "pdf",
        "starterMemos",
        "packageFiles",
        "manifestSha256",
        "notice",
    }

    for industry, details in INDUSTRIES.items():
        package_root = BUILD / industry / details["root"]
        dataset_summary_path = BUILD / industry / "dataset-summary.json"
        zip_path = PUBLIC / details["zip"]
        if not package_root.is_dir() or not zip_path.is_file() or not dataset_summary_path.is_file():
            fail(f"Missing package source or ZIP for {industry}")
        if zip_path.stat().st_size == 0:
            fail(f"Empty public ZIP: {zip_path.name}")

        dataset = json.loads(dataset_summary_path.read_text(encoding="utf-8"))
        if dataset.get("datasetVersion") != catalog["version"]:
            fail(f"Dataset version mismatch for {industry}")
        if dataset.get("generatedAt") != catalog["updatedAt"]:
            fail(f"Dataset date mismatch for {industry}")
        for key, expected in (
            ("industry", industry),
            ("shortName", details["label"]),
            ("company", details["company"]),
            ("business", details["business"]),
            ("rootName", details["root"]),
            ("notice", NOTICE),
            ("expectedDocxCount", 30),
            ("expectedPdfCount", 100),
        ):
            if dataset.get(key) != expected:
                fail(f"dataset-summary {key} mismatch for {industry}")

        current_sha = sha256(zip_path)
        item = catalog_by_industry[industry]
        if set(item) != expected_catalog_keys:
            fail(f"Unexpected catalog fields for {industry}: {sorted(item)}")
        for key, expected in (
            ("industry", industry),
            ("label", details["label"]),
            ("company", details["company"]),
            ("business", details["business"]),
            ("file", details["zip"]),
            ("url", f"/downloads/demo-data/{details['zip']}"),
            ("version", catalog["version"]),
            ("updatedAt", catalog["updatedAt"]),
            ("notice", NOTICE),
            ("workbooks", 10),
            ("sheets", details["sheets"]),
            ("dataRows", details["dataRows"]),
            ("docx", 30),
            ("pdf", 100),
            ("starterMemos", len(STARTERS)),
            ("packageFiles", 156),
        ):
            if item.get(key) != expected:
                fail(f"Catalog {key} mismatch for {industry}: {item.get(key)!r}")
        if item["bytes"] != zip_path.stat().st_size:
            fail(f"Catalog byte count mismatch: {details['zip']}")
        if current_sha != item["sha256"] or checksums.get(details["zip"]) != current_sha:
            fail(f"Public SHA-256 mismatch: {details['zip']}")

        source_hashes, _ = source_and_zip_hashes(package_root, zip_path, details["root"])
        if len(source_hashes) != item["packageFiles"]:
            fail(f"Catalog packageFiles mismatch for {industry}: {len(source_hashes)}")
        actual_manifest_sha = validate_package_manifest(package_root, source_hashes)
        if actual_manifest_sha != item["manifestSha256"]:
            fail(f"Catalog manifestSha256 mismatch: {details['zip']}")
        with zipfile.ZipFile(zip_path) as archive:
            manifest_name = f"{details['root']}/manifest.csv"
            manifest_sha = hashlib.sha256(archive.read(manifest_name)).hexdigest()
        if manifest_sha != item["manifestSha256"]:
            fail(f"Manifest SHA-256 mismatch: {details['zip']}")

        counts = validate_zip(zip_path, details["root"])
        if counts != (10, 30, 100):
            fail(f"Unexpected ZIP counts for {industry}: {counts}")

        xlsx = sorted(package_root.rglob("*.xlsx"))
        docx = sorted(package_root.rglob("*.docx"))
        pdf = sorted(package_root.rglob("*.pdf"))
        if (len(xlsx), len(docx), len(pdf)) != (10, 30, 100):
            fail(f"Unexpected source counts for {industry}")
        actual_sheets, actual_data_rows = xlsx_metrics(xlsx)
        if actual_sheets != details["sheets"] or actual_data_rows != details["dataRows"]:
            fail(
                f"Actual XLSX metrics mismatch for {industry}: "
                f"sheets={actual_sheets}, dataRows={actual_data_rows}"
            )
        if dataset.get("workbookCount") != len(xlsx):
            fail(f"dataset-summary workbookCount mismatch for {industry}")
        if dataset.get("workbookSheetCount") != actual_sheets:
            fail(f"dataset-summary workbookSheetCount mismatch for {industry}")
        if dataset.get("workbookDataRows") != actual_data_rows:
            fail(f"dataset-summary workbookDataRows mismatch for {industry}")

        workbook_reports = dataset.get("workbookReports")
        if not isinstance(workbook_reports, list) or len(workbook_reports) != len(xlsx):
            fail(f"dataset-summary workbookReports mismatch for {industry}")
        reported_workbooks = {report.get("file") for report in workbook_reports}
        actual_workbooks = {path.relative_to(package_root).as_posix() for path in xlsx}
        if reported_workbooks != actual_workbooks:
            fail(f"dataset-summary workbook paths mismatch for {industry}")
        for report in workbook_reports:
            sheets = report.get("sheets")
            if not isinstance(sheets, list) or report.get("rows") != sum(
                sheet.get("rows", -1) for sheet in sheets
            ):
                fail(f"dataset-summary sheet rows mismatch: {report.get('file')}")
            for sheet in sheets:
                preview = sheet.get("preview")
                if not isinstance(preview, str):
                    fail(f"Missing XLSX preview path: {report.get('file')}")
                expected_xlsx_previews.add(BUILD / preview)

        if not all(NOTICE in xlsx_text(path) for path in xlsx):
            fail(f"Fictional-data notice missing from XLSX in {industry}")
        if not all(NOTICE in docx_text(path) for path in docx):
            fail(f"Fictional-data notice missing from DOCX in {industry}")
        if not all(NOTICE in pdf_text(path) for path in pdf):
            fail(f"Fictional-data notice missing from PDF in {industry}")
        if not all(pdf_has_embedded_font(path) for path in pdf):
            fail(f"Embedded Japanese font missing from PDF in {industry}")

        readme_text, zip_text = zip_readable_text(zip_path)
        for required in (
            "Windows",
            "Mac",
            "すべて展開",
            "ダブルクリック",
            "コピー",
            "貼",
            "添付",
            "Work",
            "推測",
        ):
            if required not in readme_text:
                fail(f"Practical starter guidance is missing {required!r} in {details['zip']}")

        for forbidden in (
            "Windows版で進めます",
            "Windows版ChatGPT",
            "Macは対象外",
            "Macユーザーは対象外",
            "チャットへドラッグ",
            "課題に書かれたファイル名をAIへ伝える",
            "AIへファイル名を伝えて",
            "ファイル名を言うだけ",
            "AIへ「課題の",
            "毎回のアップロードは不要",
        ):
            if forbidden in zip_text:
                fail(f"Old fixed-input guidance remains in {details['zip']}: {forbidden}")

        all_docx.extend(docx)
        all_pdf.extend(pdf)
        total_workbooks += len(xlsx)
        total_sheets += actual_sheets
        total_data_rows += actual_data_rows
        summary["industries"][industry] = {
            "zip": details["zip"],
            "sha256": current_sha,
            "xlsx": len(xlsx),
            "docx": len(docx),
            "pdf": len(pdf),
            "starterMemos": len(STARTERS),
            "packageFiles": len(source_hashes),
            "sheets": actual_sheets,
            "dataRows": actual_data_rows,
        }

    renderer = find_docx_renderer()
    renderer_python = find_docx_renderer_python(renderer)
    pdftoppm = shutil.which("pdftoppm")
    pdfinfo = shutil.which("pdfinfo")
    if not pdftoppm or not pdfinfo:
        fail("pdftoppm/pdfinfo is required")

    for path in all_pdf:
        run_checked(
            [pdfinfo, str(path)],
            f"pdfinfo for {path}",
        )

    try:
        docx_workers = int(os.environ.get("AIJUKU_DOCX_RENDER_WORKERS", "1"))
    except ValueError:
        fail("AIJUKU_DOCX_RENDER_WORKERS must be an integer")
    if docx_workers < 1:
        fail("AIJUKU_DOCX_RENDER_WORKERS must be 1 or greater")
    if docx_workers == 1:
        docx_page_groups = [render_docx(path, renderer, renderer_python) for path in all_docx]
    else:
        with concurrent.futures.ThreadPoolExecutor(max_workers=docx_workers) as executor:
            docx_page_groups = list(
                executor.map(
                    lambda path: render_docx(path, renderer, renderer_python),
                    all_docx,
                )
            )
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        pdf_page_groups = list(executor.map(lambda path: render_pdf(path, pdftoppm), all_pdf))
    docx_previews = [page for pages in docx_page_groups for page in pages]
    pdf_previews = [page for pages in pdf_page_groups for page in pages]

    all_xlsx_previews = set((QA / "xlsx").rglob("*.png"))
    if all_xlsx_previews != expected_xlsx_previews:
        missing = sorted(str(path.relative_to(BUILD)) for path in expected_xlsx_previews - all_xlsx_previews)
        extra = sorted(str(path.relative_to(BUILD)) for path in all_xlsx_previews - expected_xlsx_previews)
        fail(f"XLSX preview set mismatch: missing={missing}, extra={extra}")
    for preview in sorted(all_xlsx_previews):
        if preview.stat().st_size == 0:
            fail(f"Empty XLSX preview: {preview}")
        with Image.open(preview) as image:
            image.verify()

    sheets = []
    sheets.extend(contact_sheets(sorted(all_xlsx_previews), "xlsx"))
    sheets.extend(contact_sheets(sorted(docx_previews), "docx"))
    sheets.extend(contact_sheets(sorted(pdf_previews), "pdf"))
    summary["renders"] = {
        "xlsxSheets": len(all_xlsx_previews),
        "xlsxWorkbooksValidated": total_workbooks,
        "xlsxDataRowsValidated": total_data_rows,
        "docxFiles": len(all_docx),
        "docxPages": len(docx_previews),
        "docxWorkers": docx_workers,
        "docxEmittedPdfNotices": len(all_docx),
        "pdfFiles": len(all_pdf),
        "pdfPages": len(pdf_previews),
        "contactSheets": len(sheets),
    }
    if total_sheets != len(all_xlsx_previews):
        fail(f"Validated XLSX sheet count differs from preview count: {total_sheets}")
    report = QA / "demo-data-verification.json"
    report.write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    print(f"QA report: {report}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
