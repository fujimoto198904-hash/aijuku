from __future__ import annotations

import csv
import argparse
import hashlib
import json
import os
import random
import textwrap
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


REPO_ROOT = Path(__file__).resolve().parent.parent
BUILD_ROOT = Path(
    os.environ.get("AIJUKU_DEMO_BUILD_ROOT", REPO_ROOT / "tmp/demo-data-build")
).resolve()
PUBLIC_ROOT = REPO_ROOT / "public/downloads/demo-data"
LIB_CATALOG_PATH = REPO_ROOT / "lib/demo-data-catalog.generated.json"
NOTICE = "練習用・すべて架空・外部送信禁止"
VERSION = "1.1.2"
GENERATED_AT = "2026-09-02"

def register_pdf_font() -> str:
    candidates: list[Path] = []
    if configured := os.environ.get("AIJUKU_PDF_FONT"):
        candidates.append(Path(configured))
    candidates.extend(
        sorted(
            Path("/System/Library/AssetsV2/com_apple_MobileAsset_Font8").glob(
                "*/AssetData/BIZ_UDGothic.ttc"
            )
        )
    )
    candidates.extend(
        [
            Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
            Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
            Path("C:/Windows/Fonts/msgothic.ttc"),
        ]
    )
    for candidate in candidates:
        if candidate.is_file():
            pdfmetrics.registerFont(
                TTFont("AIJUKUJapanese", str(candidate), subfontIndex=0)
            )
            return "AIJUKUJapanese"
    raise FileNotFoundError(
        "Japanese PDF font not found. Set AIJUKU_PDF_FONT to a Japanese TTF/TTC file."
    )


PDF_FONT_NAME = register_pdf_font()


@dataclass(frozen=True)
class Industry:
    key: str
    company: str
    short_name: str
    business: str
    root_name: str
    zip_base: str
    download_name: str
    accent: str
    customer_word: str
    project_word: str
    scenarios: tuple[str, ...]


INDUSTRIES = {
    "salon": Industry(
        key="salon",
        company="株式会社ひだまりデモサロン（架空）",
        short_name="ひだまり美容室",
        business="美容室2店舗の運営、予約、施術、物販、集客",
        root_name="美容室デモデータ",
        zip_base="美容室_練習用デモデータ_FULL_v1",
        download_name="salon-demo-data-full-v1.zip",
        accent="A94B3A",
        customer_word="お客様",
        project_word="店舗改善",
        scenarios=(
            "予約変更の相談",
            "カラー料金の確認",
            "来店後のフォロー",
            "キャンセルの連絡",
            "在庫不足の相談",
            "スタッフのシフト相談",
        ),
    ),
    "construction": Industry(
        key="construction",
        company="株式会社青空デモ建設（架空）",
        short_name="青空デモ建設",
        business="住宅改修、店舗改修、小規模新築、外構工事",
        root_name="建設業デモデータ",
        zip_base="建設業_練習用デモデータ_FULL_v1",
        download_name="construction-demo-data-full-v1.zip",
        accent="345FE7",
        customer_word="発注者",
        project_word="工事案件",
        scenarios=(
            "現地調査の相談",
            "見積条件の確認",
            "工程変更の相談",
            "材料納期の確認",
            "雨漏りの連絡",
            "図面の版の確認",
        ),
    ),
    "realestate": Industry(
        key="realestate",
        company="株式会社まちかどデモ不動産（架空）",
        short_name="まちかどデモ不動産",
        business="賃貸仲介、売買仲介、賃貸管理、修繕受付",
        root_name="不動産会社デモデータ",
        zip_base="不動産会社_練習用デモデータ_FULL_v1",
        download_name="real-estate-demo-data-full-v1.zip",
        accent="2E765D",
        customer_word="お客様",
        project_word="契約・管理案件",
        scenarios=(
            "物件問い合わせ",
            "内見後の相談",
            "初期費用の確認",
            "修繕受付",
            "家賃入金の確認",
            "鍵の返却確認",
        ),
    ),
}


WORD_CORE_DOCS = (
    ("01_会社とルール", "01_会社紹介_社内用.docx", "この会社のこと", "会社の特徴と仕事の全体像"),
    ("01_会社とルール", "02_組織と役割.docx", "誰が何をしているか", "部署、役割、相談先"),
    ("01_会社とルール", "03_新人向け_一日の流れ.docx", "新人の一日", "朝から退勤までの流れ"),
    ("課題", "14AIルール.docx", "AIを使う時の約束", "入れてよい情報と人が見る場所"),
    ("01_会社とルール", "05_個人情報を扱う時の手順.docx", "個人情報の扱い", "隠す、確認する、外へ出さない"),
    ("01_会社とルール", "06_経費を申請する時の流れ.docx", "経費の流れ", "領収書から確認まで"),
    ("02_仕事のやり方", "07_お客様対応の流れ.docx", "お客様対応", "受付、確認、返事、記録"),
    ("02_仕事のやり方", "08_電話対応の会話例.docx", "電話の受け方", "雑な会話から必要事項を残す"),
    ("02_仕事のやり方", "09_問い合わせ返信の文面集.docx", "返信の言い方", "短く、決めつけず、未送信で作る"),
    ("02_仕事のやり方", "10_緊急時の最初の動き.docx", "困った時の最初の動き", "止める、知らせる、記録する"),
    ("02_仕事のやり方", "11_日々の業務マニュアル.docx", "毎日の基本", "開店、確認、記録、引継ぎ"),
    ("03_ひな形", "12_契約と合意の練習用ひな形.docx", "契約の練習用見本", "法的効力のない確認用ひな形"),
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_doc_defaults(document: Document, industry: Industry) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = document.styles["Normal"]
    normal.font.name = "BIZ UDGothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string("17212B")
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, "102A36"),
        ("Heading 1", 15, industry.accent),
        ("Heading 2", 11, "102A36"),
    ):
        style = document.styles[style_name]
        style.font.name = "BIZ UDGothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(7 if style_name != "Title" else 0)
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.text = f"{NOTICE}　｜　{industry.short_name}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "BIZ UDGothic"
    header.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(industry.accent)

    footer = section.footer.paragraphs[0]
    footer.text = "教材用デモ文書｜実在の会社・人物・契約とは関係ありません"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "BIZ UDGothic"
    footer.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor.from_string("5E6A73")


def add_doc_cover(document: Document, industry: Industry, title: str, doc_id: str, subtitle: str) -> None:
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(7)
    run = kicker.add_run("DEMO DOCUMENT / 練習用")
    run.font.name = "BIZ UDGothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(industry.accent)

    title_p = document.add_paragraph(style="Title")
    title_p.add_run(title)
    subtitle_p = document.add_paragraph(subtitle)
    subtitle_p.paragraph_format.space_after = Pt(7)
    subtitle_p.runs[0].font.size = Pt(9.5)
    subtitle_p.runs[0].font.color.rgb = RGBColor.from_string("5E6A73")

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    meta_data = (
        ("文書ID", doc_id),
        ("会社", industry.company),
        ("版・作成日", f"v{VERSION} / {GENERATED_AT}"),
        ("状態", NOTICE),
    )
    for row_index, (label, value) in enumerate(meta_data):
        meta.cell(row_index, 0).text = label
        meta.cell(row_index, 1).text = value
        set_cell_shading(meta.cell(row_index, 0), industry.accent)
        for run in meta.cell(row_index, 0).paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
        for cell in meta.rows[row_index].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "BIZ UDGothic"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
                    run.font.size = Pt(9)

    warning = document.add_table(rows=1, cols=1)
    warning.alignment = WD_TABLE_ALIGNMENT.LEFT
    warning.cell(0, 0).text = "この文書は全部デモです。実在の相手への送信、契約、提出には使えません。"
    set_cell_shading(warning.cell(0, 0), "FFF0EC")
    warning.cell(0, 0).paragraphs[0].runs[0].font.bold = True
    warning.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("8A382D")


def add_simple_table(document: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]], accent: str) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
        set_cell_shading(table.cell(0, index), accent)
        for run in table.cell(0, index).paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.name = "BIZ UDGothic"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
            run.font.size = Pt(8)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = str(value)
            for run in cells[index].paragraphs[0].runs:
                run.font.name = "BIZ UDGothic"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "BIZ UDGothic")
                run.font.size = Pt(8)


def add_bullets(document: Document, values: tuple[str, ...] | list[str]) -> None:
    for value in values:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(value)


def create_core_docx(industry: Industry, spec: tuple[str, str, str, str], index: int, root: Path) -> Path:
    folder, file_name, title, subtitle = spec
    output = root / folder / file_name if folder == "課題" else root / "20_Word" / folder / file_name
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_doc_defaults(document, industry)
    doc_id = f"DOC-{industry.key.upper()}-{index:04d}"
    add_doc_cover(document, industry, title, doc_id, subtitle)

    document.add_heading("まず知っておくこと", level=1)
    opening = {
        1: f"{industry.short_name}は、{industry.business}を行う教材用の架空会社です。資料同士はIDでつながっています。",
        2: "困った時に誰へ聞けばよいか分かるよう、役割と相談先を短くまとめています。",
        3: "全部覚える必要はありません。朝、昼、終わる前の3つだけ見れば始められます。",
        4: "AIは文章や表を作る相棒です。実在の秘密を入れず、送信や判断は人が行います。",
        5: "個人情報は、必要な人だけが、必要な範囲だけ見ます。練習ではデモIDを使います。",
        6: "領収書と内容を見て、会社の決め方に沿って申請します。税務判断は人が確認します。",
    }.get(index, f"{industry.customer_word}からの連絡を受けたら、分かることと分からないことを分けて残します。")
    document.add_paragraph(opening)

    document.add_heading("この資料で使うもの", level=1)
    add_simple_table(
        document,
        ("見るもの", "書いてあること", "人が最後に見る所"),
        [
            ("元メモ", "相手が話したこと", "名前・日付・数字"),
            ("Excel", "会社の正本データ", "IDと最新版"),
            ("AIの案", "整理した文章や表", "勝手に足した内容がないか"),
        ],
        industry.accent,
    )

    document.add_heading("だいたいこの順で進める", level=1)
    add_bullets(
        document,
        (
            "短いTXTや雑メモは、中身をコピーしてチャットへ貼る（おすすめ）",
            "PDF・Word・Excelは、ファイルをチャットへ添付する",
            "一回の文章づくりはChat、実ファイルを作って何度も直す時はWorkを使う",
            "Workの入力欄でも、同じように短文を貼り、書式資料を添付できる",
            "読めない時は、短文を貼り直すか書式資料を添付し、AIに推測させない",
            "名前、日付、金額、送信先だけは人が見る",
            "完成版と、まだの所を同じフォルダへ残す",
        ),
    )

    document.add_heading("やりがちなミス", level=1)
    add_simple_table(
        document,
        ("よくあること", "その場でこうする"),
        [
            ("AIが資料にない数字を足した", "『その数字どこから？ 資料にないなら消して』と言う"),
            ("説明が長すぎる", "『大事な3つだけ。難しい言葉なしで』と言う"),
            ("もう送ったように見える", "下書き・未送信と書いて保存する"),
        ],
        industry.accent,
    )

    document.add_heading("困ったら", level=1)
    document.add_paragraph("課題コード、使ったファイル、止まった画面の3つを藤本へ見せます。秘密や実在の個人情報は隠します。")
    document.save(output)
    return output


def create_meeting_docx(industry: Industry, number: int, root: Path) -> Path:
    output = root / "20_Word/04_会議の生メモ" / f"会議の生メモ_{number:02d}.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_doc_defaults(document, industry)
    doc_id = f"DOC-{industry.key.upper()}-MTG-{number:03d}"
    add_doc_cover(document, industry, f"会議の生メモ {number:02d}", doc_id, "言い切っていない話や抜けを残した練習用メモ")
    meeting_id = f"MTG-{number:06d}"
    project = f"{industry.project_word}{number:03d}（架空）"
    document.add_heading("その場で書いたメモ", level=1)
    raw_notes = (
        f"・{project}、たぶん来週までに一回返事いる",
        "・金額は前の表のままか確認。誰が見るかはまだ",
        f"・{industry.scenarios[number % len(industry.scenarios)]}の件、相手は急いでそう",
        "・写真か資料が別のフォルダにあるはず",
        "・次の会議までに一回案を見たい。日付は後で",
    )
    add_bullets(document, raw_notes)
    document.add_heading("話したこと", level=1)
    add_simple_table(
        document,
        ("MEETING_ID", "話題", "今の状態"),
        [
            (meeting_id, "今週の進み具合", "だいたい共有した"),
            (meeting_id, "相手への返事", "文面はまだ"),
            (meeting_id, "金額と期限", "確認する人が未定"),
        ],
        industry.accent,
    )
    document.add_heading("AIに頼むなら", level=1)
    document.add_paragraph("この会議メモ、決まったこと・まだのこと・誰がいつまでにやるかに分けて。勝手に担当や期限を決めないで。")
    document.save(output)
    return output


def create_call_docx(industry: Industry, number: int, root: Path) -> Path:
    output = root / "20_Word/05_電話の生メモ" / f"電話の生メモ_{number:02d}.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_doc_defaults(document, industry)
    doc_id = f"DOC-{industry.key.upper()}-CALL-{number:03d}"
    add_doc_cover(document, industry, f"電話の生メモ {number:02d}", doc_id, industry.scenarios[(number - 1) % len(industry.scenarios)])
    call_id = f"CALL-{number:06d}"
    document.add_heading("電話の途中メモ", level=1)
    document.add_paragraph(
        f"{industry.scenarios[(number - 1) % len(industry.scenarios)]}。昨日も話したかも。相手はデモ{industry.customer_word}{number:04d}。時間は夕方なら大丈夫らしい。金額は聞かれたけど、まだ資料を見ていない。いったん確認して返すと言った。"
    )
    document.add_heading("会話の一部", level=1)
    add_simple_table(
        document,
        ("CALL_ID", "話者", "発言"),
        [
            (call_id, "相手", "あの、前に聞いた件なんですけど、ちょっと予定変えられます？"),
            (call_id, "担当", "確認します。希望の時間だけもう一度いいですか？"),
            (call_id, "相手", "夕方なら。たぶん6時くらい。難しければ別の日でも。"),
            (call_id, "担当", "分かりました。空きと金額を確認して折り返します。"),
        ],
        industry.accent,
    )
    document.add_heading("AIに頼むなら", level=1)
    document.add_paragraph("この電話、やることと予定候補と返事の下書きに分けて。決まってないことは決めないで。")
    document.save(output)
    return output


def create_handoff_docx(industry: Industry, number: int, root: Path) -> Path:
    output = root / "20_Word/06_引継ぎの走り書き" / f"担当交代の走り書き_{number:02d}.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_doc_defaults(document, industry)
    doc_id = f"DOC-{industry.key.upper()}-HANDOFF-{number:03d}"
    add_doc_cover(document, industry, f"担当交代の走り書き {number:02d}", doc_id, "次の担当が困りそうな、少し抜けのあるメモ")
    document.add_heading("前の担当が残したメモ", level=1)
    add_bullets(
        document,
        (
            f"デモ{industry.customer_word}{number:04d}、返事は早めがいい人",
            f"{industry.scenarios[number % len(industry.scenarios)]}の話が途中",
            "最新の金額はExcelにあるはず。PDFは一個古いかも",
            "次に電話する日はまだ決めてない",
            "詳しい人はデモ社員002。でも今週休みかも",
        ),
    )
    document.add_heading("AIに頼むなら", level=1)
    document.add_paragraph("この引継ぎ、次の人が最初に見るものと、まだ確認が必要なことに分けて。分からない所は空けて。")
    document.save(output)
    return output


def make_pdf_styles(industry: Industry):
    styles = getSampleStyleSheet()
    accent = colors.HexColor(f"#{industry.accent}")
    return {
        "title": ParagraphStyle(
            "JapaneseTitle",
            parent=styles["Title"],
            fontName=PDF_FONT_NAME,
            fontSize=22,
            leading=29,
            textColor=colors.HexColor("#102A36"),
            alignment=TA_LEFT,
            spaceAfter=10,
        ),
        "h1": ParagraphStyle(
            "JapaneseH1",
            parent=styles["Heading1"],
            fontName=PDF_FONT_NAME,
            fontSize=13,
            leading=18,
            textColor=accent,
            spaceBefore=8,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "JapaneseBody",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=9,
            leading=15,
            textColor=colors.HexColor("#17212B"),
            spaceAfter=5,
        ),
        "small": ParagraphStyle(
            "JapaneseSmall",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=7.5,
            leading=11,
            textColor=colors.HexColor("#5E6A73"),
        ),
        "right": ParagraphStyle(
            "JapaneseRight",
            parent=styles["BodyText"],
            fontName=PDF_FONT_NAME,
            fontSize=8,
            leading=12,
            alignment=TA_RIGHT,
            textColor=colors.HexColor("#5E6A73"),
        ),
    }


def pdf_header_footer(canvas, doc, industry: Industry, doc_id: str) -> None:
    canvas.saveState()
    width, height = doc.pagesize
    canvas.setFillColor(colors.HexColor(f"#{industry.accent}"))
    canvas.rect(0, height - 14 * mm, width, 14 * mm, fill=1, stroke=0)
    canvas.setFillColor(colors.white)
    canvas.setFont(PDF_FONT_NAME, 8)
    canvas.drawString(18 * mm, height - 9 * mm, NOTICE)
    canvas.drawRightString(width - 18 * mm, height - 9 * mm, doc_id)
    canvas.setFillColor(colors.HexColor("#5E6A73"))
    canvas.setFont(PDF_FONT_NAME, 7)
    canvas.drawString(18 * mm, 10 * mm, "教材用デモ｜実在の会社・人物・契約とは関係ありません")
    canvas.drawRightString(width - 18 * mm, 10 * mm, f"{canvas.getPageNumber()}")
    canvas.restoreState()


def make_pdf(
    output: Path,
    industry: Industry,
    title: str,
    doc_id: str,
    subtitle: str,
    rows: list[tuple[str, str]],
    notes: list[str],
    scanned: bool = False,
    landscape_page: bool = False,
) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    pagesize = landscape(A4) if landscape_page else A4
    styles = make_pdf_styles(industry)
    doc = SimpleDocTemplate(
        str(output),
        pagesize=pagesize,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=21 * mm,
        bottomMargin=17 * mm,
        title=title,
        author="藤本実学塾 デモデータ",
        subject=NOTICE,
    )
    story = [
        Spacer(1, 4 * mm),
        Paragraph("DEMO PDF / 練習用", styles["small"]),
        Paragraph(title, styles["title"]),
        Paragraph(subtitle, styles["body"]),
        Spacer(1, 2 * mm),
    ]
    meta = Table(
        [
            ["会社", industry.company],
            ["文書ID", doc_id],
            ["版・作成日", f"v{VERSION} / {GENERATED_AT}"],
            ["状態", NOTICE],
        ],
        colWidths=[38 * mm, 125 * mm if not landscape_page else 205 * mm],
    )
    meta.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{industry.accent}")),
                ("BACKGROUND", (1, 0), (1, -1), colors.HexColor("#F5F7F6")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DCE3E5")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([meta, Spacer(1, 5 * mm), Paragraph("中に書いてあること", styles["h1"])])

    table_data = [["項目", "内容（すべて架空）"]]
    table_data.extend([[Paragraph(label, styles["body"]), Paragraph(value, styles["body"])] for label, value in rows])
    table = Table(
        table_data,
        repeatRows=1,
        colWidths=[46 * mm, 117 * mm if not landscape_page else 197 * mm],
    )
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), PDF_FONT_NAME),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#102A36")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#DCE3E5")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F5F7F6")]),
            ]
        )
    )
    story.extend([table, Spacer(1, 4 * mm), Paragraph("人が見るところ", styles["h1"])])
    for note in notes:
        story.append(Paragraph(f"・{note}", styles["body"]))
    if scanned:
        story.extend(
            [
                Spacer(1, 5 * mm),
                Paragraph("OCR練習メモ", styles["h1"]),
                Paragraph("少し薄い・傾いた紙を想定した見本です。読めない所をAIに勝手に補わせず、元のExcelや人の確認へ戻ります。", styles["small"]),
            ]
        )

    doc.build(
        story,
        onFirstPage=lambda canvas, current_doc: pdf_header_footer(canvas, current_doc, industry, doc_id),
        onLaterPages=lambda canvas, current_doc: pdf_header_footer(canvas, current_doc, industry, doc_id),
    )
    return output


def build_docx_files(industry: Industry, root: Path) -> list[Path]:
    outputs: list[Path] = []
    for index, spec in enumerate(WORD_CORE_DOCS, start=1):
        outputs.append(create_core_docx(industry, spec, index, root))
    for number in range(1, 7):
        outputs.append(create_meeting_docx(industry, number, root))
    for number in range(1, 7):
        outputs.append(create_call_docx(industry, number, root))
    for number in range(1, 7):
        outputs.append(create_handoff_docx(industry, number, root))
    if len(outputs) != 30:
        raise RuntimeError(f"DOCX count mismatch for {industry.key}: {len(outputs)}")
    return outputs


def build_pdf_files(industry: Industry, root: Path) -> list[Path]:
    outputs: list[Path] = []
    randomizer = random.Random({"salon": 101, "construction": 202, "realestate": 303}[industry.key])

    outputs.append(
        make_pdf(
            root / "00_最初に見る/00_このデータは全部架空です.pdf",
            industry,
            "このデータは全部架空です",
            f"PDF-{industry.key.upper()}-0001",
            "安心して練習するために、最初に一度だけ読んでください。",
            [
                ("会社", industry.company),
                ("人物", "デモ社員・デモ顧客などの連番だけ"),
                ("連絡先", "example.invalid と 000番号だけ"),
                ("契約・税務・安全", "AIは候補整理まで。人が最終判断"),
                ("短いTXT", "中身をコピーしてチャットへ貼る（おすすめ）"),
                ("書式資料", "PDF・Word・Excelはファイルを添付する"),
                ("作業場所", "一回の文章づくりはChat、実ファイルを作って何度も直す時はWork"),
                ("Workへの入力", "短文は貼り、書式資料は添付する"),
                ("読めない時", "短文は貼り直し、書式資料は添付へ切り替え、推測させない"),
            ],
            ["実在の個人情報を混ぜない", "AIの案は送信・公開前に人が確認する", "分からない所は藤本へ画面ごと見せる"],
        )
    )

    overview_specs = (
        ("11会社.pdf", "会社案内", "この会社が何をしているかを1枚で見る"),
        ("12料金.pdf", "サービスと料金の見本", "価格は教材用の架空金額"),
        ("13組織.pdf", "組織と相談先", "困った時に誰へ聞くか"),
    )
    for index, (file_name, title, subtitle) in enumerate(overview_specs, start=2):
        outputs.append(
            make_pdf(
                root / "課題" / file_name,
                industry,
                title,
                f"PDF-{industry.key.upper()}-{index:04d}",
                subtitle,
                [
                    ("会社", industry.company),
                    ("事業", industry.business),
                    ("対象", f"架空の{industry.customer_word}と{industry.project_word}"),
                    ("連絡先", "000-0000-0000 / info@example.invalid"),
                    ("Web", "https://demo.example.invalid（無効）"),
                ],
                ["実在サイトへ公開しない", "価格・日程・約束は人が確認する"],
            )
        )

    pdf_index = 5
    for month_index in range(24):
        year = 2024 + ((month_index + 8) // 12)
        month = ((month_index + 8) % 12) + 1
        revenue = randomizer.randrange(600, 4200) * 10000
        cost = int(revenue * randomizer.uniform(0.42, 0.68))
        profit = revenue - cost
        outputs.append(
            make_pdf(
                root / "30_PDF/02_月次レポート" / f"月次レポート_{year}{month:02d}.pdf",
                industry,
                f"{year}年{month:02d}月 月次レポート",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "社長が3分で読むための教材用サンプル",
                [
                    ("売上", f"{revenue:,}円（架空）"),
                    ("費用", f"{cost:,}円（架空）"),
                    ("利益", f"{profit:,}円（架空）"),
                    ("今月よかったこと", f"{industry.scenarios[month_index % len(industry.scenarios)]}の対応が進んだ"),
                    ("まだ気になること", "期限と担当が空いている記録が一部ある"),
                ],
                ["Excelの月次シートと金額を照合する", "理由は推測だけで決めない"],
            )
        )
        pdf_index += 1

    for number in range(1, 13):
        amount = randomizer.randrange(8, 850) * 1000
        outputs.append(
            make_pdf(
                root / "30_PDF/03_見積と請求" / f"見積書_EST-{number:04d}.pdf",
                industry,
                f"見積書 EST-{number:04d}",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "比較・要約・メール作成に使う練習用見積",
                [
                    ("相手", f"デモ{industry.customer_word}{number:04d}"),
                    ("件名", f"{industry.scenarios[number % len(industry.scenarios)]}（架空）"),
                    ("税抜", f"{amount:,}円"),
                    ("税", f"{int(amount * 0.1):,}円"),
                    ("合計", f"{int(amount * 1.1):,}円"),
                    ("有効期限", "2026-09-30（教材用）"),
                ],
                ["数量と単価をExcelと比べる", "相手へ送る前に人が確認する"],
            )
        )
        pdf_index += 1
    for number in range(1, 13):
        amount = randomizer.randrange(8, 850) * 1000
        outputs.append(
            make_pdf(
                root / "30_PDF/03_見積と請求" / f"請求書_INV-{number:04d}.pdf",
                industry,
                f"請求書 INV-{number:04d}",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "入金照合に使う練習用請求書。振込先は無効です。",
                [
                    ("相手", f"デモ{industry.customer_word}{number:04d}"),
                    ("請求日", f"2026-{((number + 6) % 12) + 1:02d}-05"),
                    ("支払期限", f"2026-{((number + 7) % 12) + 1:02d}-末日"),
                    ("請求額", f"{int(amount * 1.1):,}円（架空）"),
                    ("振込先", "DEMO-BANK-0000（無効）"),
                ],
                ["請求額とExcelのINVOICE_IDを照合する", "実際の振込や督促には使わない"],
            )
        )
        pdf_index += 1

    for number in range(1, 13):
        amount = randomizer.randrange(2, 240) * 1000
        outputs.append(
            make_pdf(
                root / "30_PDF/04_仕入と領収書" / f"仕入請求書_BILL-{number:04d}.pdf",
                industry,
                f"仕入請求書 BILL-{number:04d}",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "取引先から届いた想定の教材用請求書",
                [
                    ("取引先", f"デモ取引先{number:03d}株式会社（架空）"),
                    ("発注ID", f"PO-{number:07d}"),
                    ("内容", f"デモ品目{number:03d}"),
                    ("合計", f"{int(amount * 1.1):,}円（架空）"),
                    ("状態", "照合前" if number % 4 == 0 else "照合済みデモ"),
                ],
                ["発注IDと金額をExcelで確認する", "違いがあっても勝手に修正しない"],
            )
        )
        pdf_index += 1
    for number in range(1, 13):
        amount = randomizer.randrange(1, 90) * 1000
        outputs.append(
            make_pdf(
                root / "30_PDF/04_仕入と領収書" / f"領収書_RCPT-{number:04d}.pdf",
                industry,
                f"領収書 RCPT-{number:04d}",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "OCRと経費整理に使う教材用領収書",
                [
                    ("利用日", f"2026-{((number + 5) % 12) + 1:02d}-{(number % 27) + 1:02d}"),
                    ("内容", f"デモ消耗品 {number:02d}"),
                    ("金額", f"{amount:,}円（架空）"),
                    ("支払", "会社カードDEMO" if number % 2 == 0 else "立替デモ"),
                    ("経費ID", f"EXP-{number:07d}"),
                ],
                ["日付、金額、内容をExcelと照合する", "税区分は人が確認する"],
            )
        )
        pdf_index += 1

    form_titles = (
        "問い合わせ受付票",
        "予定変更確認票",
        "経費申請票",
        "休暇申請票",
        "設備点検票",
        "事故・ヒヤリ報告",
        "会議決定メモ",
        "引継ぎ確認票",
        "公開前チェック",
        "AI利用記録",
        "見積比較表",
        "月次確認票",
    )
    for number, title in enumerate(form_titles, start=1):
        outputs.append(
            make_pdf(
                root / "30_PDF/05_帳票" / f"帳票_{number:02d}_{title}.pdf",
                industry,
                title,
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "書類整理と読み取りに使う未提出の見本",
                [
                    ("受付ID", f"FORM-{number:04d}"),
                    ("対象", f"{industry.project_word}{number:03d}（架空）"),
                    ("担当", f"デモ社員{number:03d}"),
                    ("期限", "未記入" if number % 3 == 0 else f"2026-09-{number + 1:02d}"),
                    ("状態", "下書き・未提出"),
                ],
                ["空欄を勝手に埋めない", "提出・公開は人が行う"],
            )
        )
        pdf_index += 1

    for number in range(1, 13):
        outputs.append(
            make_pdf(
                root / "90_OCR練習PDF" / f"OCR練習_{number:02d}.pdf",
                industry,
                f"少し読みづらい紙の見本 {number:02d}",
                f"PDF-{industry.key.upper()}-{pdf_index:04d}",
                "薄い文字や走り書きがある想定。答えを推測しない練習です。",
                [
                    ("受付", f"DEMO-OCR-{number:04d}"),
                    ("相手", f"デモ{industry.customer_word}{number:04d}"),
                    ("メモ", f"{industry.scenarios[number % len(industry.scenarios)]}。たぶん夕方。金額は別紙かも。"),
                    ("数字", f"{randomizer.randrange(10, 990):,}00円？（要確認）"),
                    ("状態", "読み取り前"),
                ],
                ["読めない文字を勝手に確定しない", "元のExcel、別紙、人の確認へ戻る"],
                scanned=True,
                landscape_page=number % 4 == 0,
            )
        )
        pdf_index += 1

    if len(outputs) != 100:
        raise RuntimeError(f"PDF count mismatch for {industry.key}: {len(outputs)}")
    return outputs


def sha256_file(file_path: Path) -> str:
    digest = hashlib.sha256()
    with file_path.open("rb") as source:
        for chunk in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def collect_files(root: Path) -> list[Path]:
    return sorted(
        file_path
        for file_path in root.rglob("*")
        if file_path.is_file() and file_path.name not in {"SHA256SUMS.txt", "manifest.csv"}
    )


def write_package_docs(industry: Industry, root: Path, summary: dict) -> None:
    readme = f"""{industry.short_name} 練習用デモデータ

{NOTICE}

最初に、素材の渡し方と作業場所を別々に選びます。

1. ZIPを展開する（Windowsは右クリックして「すべて展開」、Macはダブルクリック）
2. 短いTXTや雑メモは、中身をコピーしてチャットへ貼る（おすすめ）
3. PDF・Word・Excelは、ファイルをチャットへ添付する
4. 一回の文章づくりはChat、実ファイルを作って何度も直す時はChatGPTデスクトップのWorkを使う。Workの入力欄でも、同じように貼付・添付できます

まず使う場所

- 00_最初に見る/00_ファイル早見表.xlsx
- 課題/01メール.txt〜10AI安全.txt（短いメモ。中身をコピーして貼れます）
- 課題/11会社.pdf・12料金.pdf・13組織.pdf・14AIルール.docx（書式ごと見たい時は添付します）

会社まるごとの資料は、課題に書かれた時だけ開けば大丈夫です。
読めない時は、短文ならコピーして貼り直し、PDF・Word・Excelならファイル添付へ切り替えます。読めない所をAIに推測させません。

入っている量

- Excel: {summary['workbookCount']}ファイル / {summary['workbookSheetCount']}シート / {summary['workbookDataRows']:,}データ行
- Word: 30ファイル
- PDF: 100ファイル
- 雑な開始メモ: 10ファイル

大事なこと

- 全部架空です。実在データを混ぜないでください。
- AIが作った文章、金額、日付、予定は、人が最後に見ます。
- 契約、税務、法務、安全、送信、公開は人が決めます。
- 分からなければ、課題コード・使ったファイル・止まった画面を藤本へ見せます。
"""
    (root / "README_最初に読んでください.txt").write_text(readme, encoding="utf-8")

    completed_folder = root / "完成"
    completed_folder.mkdir(parents=True, exist_ok=True)
    (completed_folder / "ここに完成品を入れます.txt").write_text(
        "このフォルダは自分の完成品置き場です。\n"
        "ChatGPTが作ったExcel、Word、PDF、画像、サイト、アプリの作業版をここへ残します。\n"
        "名前、数字、期限、送信先、公開状態は人が確認してください。\n",
        encoding="utf-8",
    )

    relationships = f"""# データのつながり

> {NOTICE}

```text
CUS（{industry.customer_word}）
  ├─ INQ（問い合わせ）→ CALL（電話）→ TODO（やること）→ EVT（予定候補）
  └─ {industry.key.upper()}-PRJ（{industry.project_word}）→ SAL（売上）→ INV（請求）

VEN（取引先）→ PO（発注）→ POL（発注明細）→ MOV（在庫の動き）

MTG（会議）→ DEC（決まったこと）→ TSK（作業）
```

表をつなぐ時は、名前ではなくIDを使います。空欄や不一致は練習用に少し残しています。勝手に消したりまとめたりせず、人へ確認する候補として扱います。
"""
    (root / "99_データ辞書/IDのつながり.md").write_text(relationships, encoding="utf-8")

    dictionary = [
        ("EMPLOYEE_ID", "社員", "EMP-0001", "01_会社_社員_勤怠.xlsx"),
        ("CUSTOMER_ID", industry.customer_word, "CUS-000001", "02_顧客_営業.xlsx"),
        ("VENDOR_ID", "取引先", "VEN-0001", "04_仕入_取引先_在庫.xlsx"),
        ("PROJECT_ID", industry.project_word, f"{industry.key.upper()}-PRJ-00001", "07_会議_案件_進捗.xlsx"),
        ("INQUIRY_ID", "問い合わせ", "INQ-0000001", "06_問い合わせ_電話_ToDo.xlsx"),
        ("SALES_ID", "売上", "SAL-0000001", "03_売上_請求_入金.xlsx"),
        ("INVOICE_ID", "請求", "INV-0000001", "03_売上_請求_入金.xlsx"),
        ("PURCHASE_ORDER_ID", "発注", "PO-0000001", "04_仕入_取引先_在庫.xlsx"),
    ]
    with (root / "99_データ辞書/データ辞書.csv").open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.writer(output)
        writer.writerow(("列名", "意味", "例", "主なファイル"))
        writer.writerows(dictionary)


def write_manifest_and_hashes(root: Path) -> tuple[list[dict], str]:
    files = collect_files(root)
    rows: list[dict] = []
    for file_path in files:
        relative = file_path.relative_to(root).as_posix()
        rows.append(
            {
                "path": relative,
                "format": file_path.suffix.lower().lstrip("."),
                "bytes": file_path.stat().st_size,
                "sha256": sha256_file(file_path),
                "notice": NOTICE,
            }
        )
    manifest_path = root / "manifest.csv"
    with manifest_path.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.DictWriter(output, fieldnames=("path", "format", "bytes", "sha256", "notice"))
        writer.writeheader()
        writer.writerows(rows)
    rows.append(
        {
            "path": "manifest.csv",
            "format": "csv",
            "bytes": manifest_path.stat().st_size,
            "sha256": sha256_file(manifest_path),
            "notice": NOTICE,
        }
    )
    hash_path = root / "SHA256SUMS.txt"
    hash_path.write_text(
        "".join(f"{row['sha256']}  {row['path']}\n" for row in rows),
        encoding="utf-8",
    )
    return rows, sha256_file(manifest_path)


def make_zip(root: Path, destination: Path) -> None:
    destination.parent.mkdir(parents=True, exist_ok=True)
    temporary = destination.with_name(f".{destination.name}.tmp")
    temporary.unlink(missing_ok=True)
    try:
        with zipfile.ZipFile(
            temporary,
            "w",
            compression=zipfile.ZIP_DEFLATED,
            compresslevel=6,
        ) as archive:
            for file_path in sorted(root.rglob("*")):
                if file_path.is_file():
                    archive.write(
                        file_path,
                        arcname=(Path(root.name) / file_path.relative_to(root)).as_posix(),
                    )
        with zipfile.ZipFile(temporary) as archive:
            if archive.testzip() is not None:
                raise RuntimeError(f"Generated ZIP failed CRC validation: {destination.name}")
        temporary.replace(destination)
    finally:
        temporary.unlink(missing_ok=True)


def build_industry(industry: Industry) -> dict:
    industry_build = BUILD_ROOT / industry.key
    summary_path = industry_build / "dataset-summary.json"
    if not summary_path.exists():
        raise FileNotFoundError(f"Run build_demo_data.mjs first: {summary_path}")
    summary = json.loads(summary_path.read_text(encoding="utf-8"))
    if summary.get("datasetVersion") != VERSION or summary.get("generatedAt") != GENERATED_AT:
        raise RuntimeError(
            f"Version/date mismatch in {summary_path}. "
            "Update both generators and rebuild all three industries together."
        )
    if summary.get("industry") != industry.key or summary.get("rootName") != industry.root_name:
        raise RuntimeError(f"Industry/root mismatch in {summary_path}")
    root = industry_build / industry.root_name
    for sidecar in root.rglob("*.inspect.ndjson"):
        sidecar.unlink()
    docx_outputs = build_docx_files(industry, root)
    pdf_outputs = build_pdf_files(industry, root)
    write_package_docs(industry, root, summary)
    manifest_rows, manifest_hash = write_manifest_and_hashes(root)
    destination = PUBLIC_ROOT / industry.download_name
    make_zip(root, destination)
    return {
        "industry": industry.key,
        "label": industry.short_name,
        "company": industry.company,
        "business": industry.business,
        "file": industry.download_name,
        "url": f"/downloads/demo-data/{industry.download_name}",
        "bytes": destination.stat().st_size,
        "sha256": sha256_file(destination),
        "version": VERSION,
        "updatedAt": GENERATED_AT,
        "workbooks": summary["workbookCount"],
        "sheets": summary["workbookSheetCount"],
        "dataRows": summary["workbookDataRows"],
        "docx": len(docx_outputs),
        "pdf": len(pdf_outputs),
        "starterMemos": len(summary["starterFiles"]),
        "packageFiles": len(manifest_rows) + 1,
        "manifestSha256": manifest_hash,
        "notice": NOTICE,
    }


def validate_complete_public_catalog(catalog: list[dict]) -> None:
    expected_industries = set(INDUSTRIES)
    actual_industries = [item.get("industry") for item in catalog]
    if len(actual_industries) != len(set(actual_industries)):
        raise RuntimeError(f"Duplicate industries in public catalog: {actual_industries}")
    if set(actual_industries) != expected_industries:
        raise RuntimeError(
            "Public catalog must contain salon, construction, and realestate together. "
            f"Found: {actual_industries}"
        )
    for item in catalog:
        industry = INDUSTRIES[item["industry"]]
        if item.get("version") != VERSION or item.get("updatedAt") != GENERATED_AT:
            raise RuntimeError(
                "A one-industry rebuild would mix versions or dates. "
                "Run both generators without --industry to rebuild all three packages."
            )
        if item.get("file") != industry.download_name:
            raise RuntimeError(f"Unexpected public filename for {industry.key}: {item.get('file')}")
        zip_path = PUBLIC_ROOT / industry.download_name
        if not zip_path.is_file() or zip_path.stat().st_size == 0:
            raise RuntimeError(f"Missing or empty public ZIP: {zip_path}")
        if item.get("bytes") != zip_path.stat().st_size:
            raise RuntimeError(f"Public ZIP byte count mismatch: {zip_path.name}")
        if item.get("sha256") != sha256_file(zip_path):
            raise RuntimeError(f"Public ZIP SHA-256 mismatch: {zip_path.name}")


def partial_build_baseline() -> list[dict]:
    catalog_path = PUBLIC_ROOT / "catalog.json"
    if not catalog_path.is_file():
        raise RuntimeError(
            "--industry cannot create a partial public catalog. "
            "Run both generators without --industry first."
        )
    payload = json.loads(catalog_path.read_text(encoding="utf-8"))
    if payload.get("version") != VERSION or payload.get("updatedAt") != GENERATED_AT:
        raise RuntimeError(
            "--industry is blocked because the existing public catalog has a different "
            "version/date. Rebuild all three industries together."
        )
    current = payload.get("packages")
    if not isinstance(current, list):
        raise RuntimeError("Invalid packages in the existing public catalog")
    validate_complete_public_catalog(current)
    return current


def write_public_catalog(catalog: list[dict]) -> None:
    validate_complete_public_catalog(catalog)
    catalog = sorted(catalog, key=lambda item: ("salon", "construction", "realestate").index(item["industry"]))
    PUBLIC_ROOT.mkdir(parents=True, exist_ok=True)
    catalog_text = (
        json.dumps(
            {"version": VERSION, "updatedAt": GENERATED_AT, "packages": catalog},
            ensure_ascii=False,
            indent=2,
        )
        + "\n"
    )
    (PUBLIC_ROOT / "catalog.json").write_text(catalog_text, encoding="utf-8")
    LIB_CATALOG_PATH.write_text(catalog_text, encoding="utf-8")
    (PUBLIC_ROOT / "SHA256SUMS.txt").write_text(
        "".join(f"{item['sha256']}  {item['file']}\n" for item in catalog),
        encoding="utf-8",
    )
    (BUILD_ROOT / "office-build-manifest.json").write_text(
        json.dumps(catalog, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(catalog, ensure_ascii=False, indent=2))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--industry", choices=tuple(INDUSTRIES), help="1業種だけ生成する")
    args = parser.parse_args()
    PUBLIC_ROOT.mkdir(parents=True, exist_ok=True)

    if args.industry:
        current = partial_build_baseline()
        built = build_industry(INDUSTRIES[args.industry])
        current = [item for item in current if item.get("industry") != args.industry]
        current.append(built)
        write_public_catalog(current)
        return

    write_public_catalog([build_industry(industry) for industry in INDUSTRIES.values()])


if __name__ == "__main__":
    main()
